from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone

from fastapi import APIRouter, File, Form, HTTPException, UploadFile

from .. import db
from ..ai import store
from ..schemas import Document
from .matters import log_activity

router = APIRouter(prefix="/api/v1/documents", tags=["documents"])


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _extract_text(filename: str, data: bytes) -> tuple[str, int]:
    name = filename.lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [p.extract_text() or "" for p in reader.pages]
        return "\n".join(pages), len(pages)
    if name.endswith(".docx"):
        from docx import Document as Docx

        doc = Docx(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        return text, max(1, text.count("\n") // 40)
    # txt / md / anything else -> best-effort decode
    text = data.decode("utf-8", errors="ignore")
    return text, max(1, text.count("\n") // 40)


@router.get("", response_model=list[Document])
def list_documents(matter_id: str) -> list[Document]:
    with db.connect() as conn:
        rows = conn.execute(
            "SELECT id, matter_id, name, kind, status, pages, created_at "
            "FROM documents WHERE matter_id = ? ORDER BY created_at DESC",
            (matter_id,),
        ).fetchall()
    return [Document(**r) for r in rows]


@router.post("/upload", response_model=Document, status_code=201)
async def upload_document(
    matter_id: str = Form(...),
    file: UploadFile = File(...),
) -> Document:
    with db.connect() as conn:
        exists = conn.execute("SELECT 1 FROM matters WHERE id = ?", (matter_id,)).fetchone()
    if not exists:
        raise HTTPException(404, "Matter not found")

    data = await file.read()
    text, pages = _extract_text(file.filename or "document.txt", data)
    if not text.strip():
        raise HTTPException(422, "Could not extract text from the uploaded file")

    doc_id = str(uuid.uuid4())
    kind = (file.filename or "").split(".")[-1].lower()
    now = _now()
    with db.connect() as conn:
        conn.execute(
            "INSERT INTO documents (id, matter_id, name, kind, status, pages, text, created_at) "
            "VALUES (?, ?, ?, ?, 'indexing', ?, ?, ?)",
            (doc_id, matter_id, file.filename, kind, pages, text, now),
        )

    # Chunk + embed (Voyage). Synchronous for MVP; move to a worker for large corpora.
    n_chunks = store.index_document(doc_id, matter_id, text)

    with db.connect() as conn:
        conn.execute("UPDATE documents SET status = 'ready' WHERE id = ?", (doc_id,))

    log_activity(matter_id, "document", "Document ingested", f"{file.filename} · {n_chunks} chunks")
    return Document(
        id=doc_id, matter_id=matter_id, name=file.filename or "document",
        kind=kind, status="ready", pages=pages, created_at=now,
    )
